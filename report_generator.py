"""
report_generator.py
---------------------
Generates a downloadable PDF and DOCX report summarizing:
- the application (tabular data)
- the text-statement analysis
- the bank-statement (PDF) analysis
- the ID document (image) quality check
- the ANN's prediction, confidence, and explanation
- the human reviewer's final decision (approve/reject/modify + notes)

Satisfies the "Generate a downloadable PDF or DOCX report" requirement --
we provide BOTH formats.
"""

import io
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
)

from docx import Document
from docx.shared import Pt, RGBColor


def _rows(app_data, text_result, pdf_result, image_result, prediction, decision):
    return [
        ("Applicant / Application", [
            ("Loan ID", app_data.get("loan_id", "N/A")),
            ("Gender", app_data.get("Gender")),
            ("Married", app_data.get("Married")),
            ("Dependents", app_data.get("Dependents")),
            ("Education", app_data.get("Education")),
            ("Self Employed", app_data.get("Self_Employed")),
            ("Applicant Income", app_data.get("ApplicantIncome")),
            ("Co-applicant Income", app_data.get("CoapplicantIncome")),
            ("Loan Amount (thousands)", app_data.get("LoanAmount")),
            ("Loan Term (months)", app_data.get("Loan_Amount_Term")),
            ("Credit History", app_data.get("Credit_History")),
            ("Property Area", app_data.get("Property_Area")),
        ]),
        ("Text Statement Analysis (LLM reasoning)", [
            ("Summary", text_result.get("summary", "N/A")),
            ("Flags", ", ".join(text_result.get("flags", [])) or "None"),
            ("Engine", text_result.get("engine", "N/A")),
        ]),
        ("Bank Statement Analysis (PDF)", [
            ("Net cash flow", f"{pdf_result.get('net_flow', 0):,.0f}"),
            ("Overdraft/NSF mentions", pdf_result.get("overdraft_mentions", 0)),
            ("Risk flag", pdf_result.get("risk_flag", "N/A")),
        ]),
        ("ID Document Check (Image)", [
            ("Quality", image_result.get("quality", "N/A")),
            ("Issues", ", ".join(image_result.get("issues", [])) or "None"),
            ("Resolution", f"{image_result.get('width','?')}x{image_result.get('height','?')}"),
        ]),
        ("AI Recommendation (ANN + XAI)", [
            ("Prediction", prediction.get("label", "N/A")),
            ("Confidence", f"{prediction.get('confidence', 0)*100:.1f}%"),
            ("Reasoning", prediction.get("reasoning", "N/A")),
        ]),
        ("Human-in-the-Loop Decision", [
            ("Final decision", decision.get("action", "Pending")),
            ("Reviewer notes", decision.get("notes", "")),
            ("Reviewer", decision.get("reviewer", "N/A")),
            ("Timestamp", decision.get("timestamp", datetime.utcnow().isoformat())),
        ]),
    ]


def build_pdf_report(app_data, text_result, pdf_result, image_result, prediction, decision) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleX", parent=styles["Title"], fontSize=18)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], spaceBefore=12, textColor=colors.HexColor("#1f3a5f"))
    body = styles["BodyText"]

    story = [
        Paragraph("Bank Loan Assistant — Decision Report", title_style),
        Paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", body),
        Spacer(1, 12),
    ]

    for section_title, rows in _rows(app_data, text_result, pdf_result, image_result, prediction, decision):
        story.append(Paragraph(section_title, h2))
        table_data = [[Paragraph(str(k), body), Paragraph(str(v), body)] for k, v in rows]
        t = Table(table_data, colWidths=[5.5 * cm, 10.5 * cm])
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f4f8")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
        ]))
        story.append(t)
        story.append(Spacer(1, 8))

    doc.build(story)
    return buf.getvalue()


def build_docx_report(app_data, text_result, pdf_result, image_result, prediction, decision) -> bytes:
    doc = Document()

    title = doc.add_heading("Bank Loan Assistant — Decision Report", level=0)
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")

    for section_title, rows in _rows(app_data, text_result, pdf_result, image_result, prediction, decision):
        doc.add_heading(section_title, level=1)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for k, v in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = str(k)
            row_cells[1].text = str(v)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
